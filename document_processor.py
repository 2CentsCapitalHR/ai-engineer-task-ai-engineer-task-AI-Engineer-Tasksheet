import io
import re
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.parser import OxmlElement
from docx.oxml.ns import qn

class DocumentProcessor:
    """Handles document parsing, type identification, and markup."""
    
    def __init__(self):
        self.document_patterns = {
            'Articles of Association': [
                r'articles?\s+of\s+association',
                r'memorandum\s+and\s+articles',
                r'company\s+constitution'
            ],
            'Memorandum of Association': [
                r'memorandum\s+of\s+association',
                r'memorandum\s+of\s+understanding',
                r'company\s+objects'
            ],
            'Board Resolution': [
                r'board\s+resolution',
                r'directors?\s+resolution',
                r'board\s+meeting\s+minutes'
            ],
            'Shareholder Resolution': [
                r'shareholder\s+resolution',
                r'members?\s+resolution',
                r'general\s+meeting'
            ],
            'Incorporation Application': [
                r'incorporation\s+application',
                r'company\s+registration',
                r'certificate\s+of\s+incorporation'
            ],
            'UBO Declaration': [
                r'ultimate\s+beneficial\s+owner',
                r'ubo\s+declaration',
                r'beneficial\s+ownership'
            ],
            'Register of Members': [
                r'register\s+of\s+members',
                r'register\s+of\s+directors',
                r'member\s+register'
            ],
            'Employment Contract': [
                r'employment\s+contract',
                r'employment\s+agreement',
                r'job\s+contract'
            ],
            'Commercial Agreement': [
                r'commercial\s+agreement',
                r'service\s+agreement',
                r'supply\s+agreement'
            ],
            'Compliance Policy': [
                r'compliance\s+policy',
                r'risk\s+policy',
                r'governance\s+policy'
            ]
        }
    
    def parse_document(self, uploaded_file):
        """Parse a .docx file and extract text content."""
        try:
            # Read the uploaded file
            file_content = uploaded_file.read()
            uploaded_file.seek(0)  # Reset file pointer
            
            # Parse with python-docx
            doc = Document(io.BytesIO(file_content))
            
            # Extract text content
            content = {
                'paragraphs': [],
                'tables': [],
                'headers': [],
                'footers': []
            }
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    content['paragraphs'].append({
                        'text': para.text,
                        'style': para.style.name if para.style else 'Normal'
                    })
            
            # Extract tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                content['tables'].append(table_data)
            
            # Extract headers and footers
            for section in doc.sections:
                if section.header:
                    for para in section.header.paragraphs:
                        if para.text.strip():
                            content['headers'].append(para.text)
                
                if section.footer:
                    for para in section.footer.paragraphs:
                        if para.text.strip():
                            content['footers'].append(para.text)
            
            return content
            
        except Exception as e:
            raise Exception(f"Error parsing document: {str(e)}")
    
    def identify_document_type(self, content):
        """Identify the type of document based on content patterns."""
        # Combine all text content
        all_text = ""
        for para in content['paragraphs']:
            all_text += para['text'] + " "
        for header in content['headers']:
            all_text += header + " "
        
        all_text = all_text.lower()
        
        # Check patterns for each document type
        for doc_type, patterns in self.document_patterns.items():
            for pattern in patterns:
                if re.search(pattern, all_text, re.IGNORECASE):
                    return doc_type
        
        return "Unknown Document Type"
    
    def create_marked_document(self, uploaded_file, issues):
        """Create a marked-up version of the document with comments."""
        try:
            # Read the uploaded file again
            uploaded_file.seek(0)
            file_content = uploaded_file.read()
            uploaded_file.seek(0)
            
            # Load document
            doc = Document(io.BytesIO(file_content))
            
            # Add comments for each issue
            for issue in issues:
                self._add_comment_to_document(doc, issue)
            
            # Add summary comment at the beginning
            self._add_summary_comment(doc, issues)
            
            # Save to bytes
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            
            return output.getvalue()
            
        except Exception as e:
            # Reset file pointer in case of error
            uploaded_file.seek(0)
            raise Exception(f"Error creating marked document: {str(e)}")
    
    def _add_comment_to_document(self, doc, issue):
        """Add a comment to the document for a specific issue."""
        # Find the relevant paragraph/section
        target_text = issue.get('section', '').lower()
        
        for para in doc.paragraphs:
            if target_text and target_text in para.text.lower():
                # Highlight the paragraph
                for run in para.runs:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                
                # Add a new paragraph with the issue description
                issue_para = doc.add_paragraph()
                issue_run = issue_para.add_run(f"[ADGM REVIEW] {issue['severity']} Issue: {issue['issue']}")
                issue_run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
                issue_run.bold = True
                
                if issue.get('suggestion'):
                    suggestion_para = doc.add_paragraph()
                    suggestion_run = suggestion_para.add_run(f"Suggestion: {issue['suggestion']}")
                    suggestion_run.font.color.rgb = RGBColor(0, 100, 0)  # Green color
                
                if issue.get('adgm_reference'):
                    ref_para = doc.add_paragraph()
                    ref_run = ref_para.add_run(f"ADGM Reference: {issue['adgm_reference']}")
                    ref_run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color
                    ref_run.italic = True
                
                break
    
    def _add_summary_comment(self, doc, issues):
        """Add a summary comment at the beginning of the document."""
        # Insert at the beginning
        summary_para = doc.paragraphs[0].insert_paragraph_before()
        summary_run = summary_para.add_run("=== ADGM COMPLIANCE REVIEW SUMMARY ===")
        summary_run.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue
        summary_run.bold = True
        summary_run.font.size = summary_run.font.size * 1.2 if summary_run.font.size else None
        
        # Add issue count
        high_issues = len([i for i in issues if i['severity'] == 'High'])
        medium_issues = len([i for i in issues if i['severity'] == 'Medium'])
        low_issues = len([i for i in issues if i['severity'] == 'Low'])
        
        count_para = doc.paragraphs[0].insert_paragraph_before()
        count_text = f"Total Issues Found: {len(issues)} (High: {high_issues}, Medium: {medium_issues}, Low: {low_issues})"
        count_run = count_para.add_run(count_text)
        count_run.font.color.rgb = RGBColor(139, 0, 0)  # Dark red
        
        # Add separator
        separator_para = doc.paragraphs[0].insert_paragraph_before()
        separator_run = separator_para.add_run("=" * 50)
        separator_run.font.color.rgb = RGBColor(0, 0, 139)
