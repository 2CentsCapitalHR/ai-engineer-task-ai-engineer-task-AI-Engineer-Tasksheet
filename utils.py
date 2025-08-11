# utils.py

from docx import Document
from docx.shared import RGBColor
from doc_checklist import ADGM_CHECKLISTS

def extract_text(doc_path):
    """Extract all text from a .docx file."""
    doc = Document(doc_path)
    return "\n".join([p.text for p in doc.paragraphs])

def detect_document_type(text):
    """
    Basic document type detection based on keywords.
    You can upgrade this later using an AI classifier.
    """
    if "Articles of Association" in text:
        return "Articles of Association"
    if "Memorandum of Association" in text:
        return "Memorandum of Association"
    if "Resolution" in text:
        return "Board Resolution Templates"
    return "Unknown"

def compare_with_checklist(process, uploaded_types):
    """
    Compare detected document types with the required ADGM checklist.
    Returns (required_docs_list, missing_docs_list)
    """
    required = set(ADGM_CHECKLISTS.get(process, []))
    uploaded = set(uploaded_types)
    missing = list(required - uploaded)
    return list(required), missing

def insert_comment(doc_path, output_path, comments):
    """
    Adds comments into the docx file as red text inline.
    comments: dict where key=paragraph index, value=comment text
    """
    doc = Document(doc_path)
    for idx, text in comments.items():
        if idx < len(doc.paragraphs):
            run = doc.paragraphs[idx].add_run(f" [COMMENT: {text}]")
            run.font.color.rgb = RGBColor(255, 0, 0)  # red text
    doc.save(output_path)
