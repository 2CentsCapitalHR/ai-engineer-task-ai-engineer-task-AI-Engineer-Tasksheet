"""Tests for document parser functionality."""

import pytest
import tempfile
import os
from pathlib import Path
from docx import Document

# Add src to path for testing
import sys
sys.path.append(str(Path(__file__).parent.parent / "src"))

from src.core.document_parser import DocumentParser
from src.models import DocumentType


class TestDocumentParser:
    """Test cases for DocumentParser class."""
    
    def setup_method(self):
        """Set up test fixtures."""
        self.parser = DocumentParser()
    
    def create_test_docx(self, content: str, filename: str = "test.docx") -> str:
        """Create a test DOCX file with given content."""
        doc = Document()
        doc.add_paragraph(content)
        
        temp_dir = tempfile.mkdtemp()
        file_path = os.path.join(temp_dir, filename)
        doc.save(file_path)
        
        return file_path
    
    def test_parse_articles_of_association(self):
        """Test parsing Articles of Association document."""
        content = """
        ARTICLES OF ASSOCIATION
        
        1. Company Name
        The name of the company is Test Company Limited.
        
        2. Share Capital
        The authorized share capital of the company is AED 100,000.
        
        3. Directors Powers
        The directors shall have the power to manage the company.
        
        4. General Meetings
        Annual general meetings shall be held each year.
        """
        
        file_path = self.create_test_docx(content)
        
        try:
            result = self.parser.parse_document(file_path)
            
            assert result['document_type'] == DocumentType.ARTICLES_OF_ASSOCIATION
            assert result['type_confidence'] > 0.3
            assert result['word_count'] > 0
            assert 'share capital' in result['text_content'].lower()
            assert 'directors' in result['text_content'].lower()
            
        finally:
            os.unlink(file_path)
    
    def test_parse_memorandum_of_association(self):
        """Test parsing Memorandum of Association document."""
        content = """
        MEMORANDUM OF ASSOCIATION
        
        1. Company Name
        The name of the company is Test Company Limited.
        
        2. Objects of the Company
        The objects of the company are to carry on business activities.
        
        3. Liability of Members
        The liability of members is limited by shares.
        
        4. Authorized Share Capital
        The authorized share capital is AED 100,000.
        """
        
        file_path = self.create_test_docx(content)
        
        try:
            result = self.parser.parse_document(file_path)
            
            assert result['document_type'] == DocumentType.MEMORANDUM_OF_ASSOCIATION
            assert result['type_confidence'] > 0.3
            assert 'objects of the company' in result['text_content'].lower()
            assert 'liability of members' in result['text_content'].lower()
            
        finally:
            os.unlink(file_path)
    
    def test_parse_board_resolution(self):
        """Test parsing Board Resolution document."""
        content = """
        BOARD RESOLUTION
        
        At a meeting of the Board of Directors held on [DATE].
        
        RESOLVED THAT:
        
        1. The company shall open a bank account.
        2. The directors are authorized to sign documents.
        
        Quorum present: 3 directors
        """
        
        file_path = self.create_test_docx(content)
        
        try:
            result = self.parser.parse_document(file_path)
            
            assert result['document_type'] == DocumentType.BOARD_RESOLUTION
            assert result['type_confidence'] > 0.3
            assert 'resolved that' in result['text_content'].lower()
            assert 'board of directors' in result['text_content'].lower()
            
        finally:
            os.unlink(file_path)
    
    def test_parse_unknown_document(self):
        """Test parsing unknown document type."""
        content = """
        This is a random document with no specific legal structure.
        It contains general text without any legal keywords.
        """
        
        file_path = self.create_test_docx(content)
        
        try:
            result = self.parser.parse_document(file_path)
            
            assert result['document_type'] == DocumentType.OTHER
            assert result['type_confidence'] == 0.0
            
        finally:
            os.unlink(file_path)
    
    def test_extract_structured_content(self):
        """Test extraction of structured content."""
        content = """
        1. First Section
        This is the first section.
        
        1.1 First subsection
        This is a subsection.
        
        2. Second Section
        This is the second section.
        
        Signature: _______________
        Date: _______________
        """
        
        file_path = self.create_test_docx(content)
        
        try:
            result = self.parser.parse_document(file_path)
            structured = result['structured_content']
            
            assert len(structured['sections']) > 0
            assert len(structured['clauses']) > 0
            assert len(structured['signatures']) > 0
            
        finally:
            os.unlink(file_path)
    
    def test_add_comment_to_document(self):
        """Test adding comments to document."""
        content = "This is a test paragraph for commenting."
        file_path = self.create_test_docx(content)
        
        try:
            doc = Document(file_path)
            original_paragraphs = len(doc.paragraphs)
            
            # Add comment to first paragraph
            updated_doc = self.parser.add_comment_to_document(
                doc, 0, "This is a test comment", "test"
            )
            
            # Should have more paragraphs after adding comment
            assert len(updated_doc.paragraphs) > original_paragraphs
            
        finally:
            os.unlink(file_path)


if __name__ == "__main__":
    pytest.main([__file__])
