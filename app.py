import streamlit as st
from io import BytesIO
import docx
import faiss
import pickle
from sentence_transformers import SentenceTransformer
import ollama
import json
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# -----------------------------
# PREDEFINED DOCUMENT CHECKLIST
# -----------------------------
ADGM_CHECKLISTS = {
    "Company Incorporation": [
        "Articles of Association",
        "Memorandum of Association",
        "Board Resolution",
        "Incorporation Application Form",
        "Register of Members and Directors"
    ]
}

# -----------------------------
# LOAD FAISS INDEX AND METADATA
# -----------------------------
INDEX_FILE = "adgm_faiss.index"
META_FILE = "adgm_metadata.pkl"

@st.cache_resource
def load_vector_store():
    index = faiss.read_index(INDEX_FILE)
    with open(META_FILE, "rb") as f:
        data = pickle.load(f)
    embedding_model = SentenceTransformer("all-MiniLM-L6-v2")  # must match indexing model
    return index, data["docs"], data["metadata"], embedding_model

index, docs, metadata, embedding_model = load_vector_store()

# -----------------------------
# SIMPLE DOCUMENT TYPE DETECTOR
# -----------------------------
def detect_document_type(doc):
    full_text = "\n".join([p.text for p in doc.paragraphs])
    full_text_lower = full_text.lower()

    if "articles of association" in full_text_lower:
        return "Articles of Association"
    elif "memorandum of association" in full_text_lower or "mou" in full_text_lower:
        return "Memorandum of Association"
    elif "board resolution" in full_text_lower:
        return "Board Resolution"
    elif "incorporation application" in full_text_lower:
        return "Incorporation Application Form"
    elif "register of members" in full_text_lower:
        return "Register of Members and Directors"
    else:
        return "Unknown Document"

def retrieve_adgm_references(query, top_k=3):
    query_vec = embedding_model.encode([query])
    distances, indices = index.search(query_vec, top_k)
    results = []
    for idx, dist in zip(indices[0], distances[0]):
        results.append({
            "source": metadata[idx]["source"],
            "score": float(dist),
            "excerpt": docs[idx][:300] + "..."
        })
    return results

def analyze_with_llm(doc_text, references):
    context_text = "\n\n".join([f"Source: {ref['source']}\nExcerpt: {ref['excerpt']}" for ref in references])
    prompt = f"""
You are an ADGM Corporate Compliance Assistant.
You will review the following document text for ADGM compliance issues.
Use the retrieved ADGM law excerpts for accuracy.

ADGM Legal References:
{context_text}

Document Text:
{doc_text}

Instructions:
- Identify compliance issues or red flags.
- Cite relevant ADGM rules using the provided excerpts where possible.
- Suggest clause improvements if applicable.
- Return structured JSON with:
  document_name, issues_found (with section/summary/severity/suggestion)
"""
    response = ollama.chat(model="mistral", messages=[{"role": "user", "content": prompt}])
    return response['message']['content']

def parse_llm_json(llm_output):
    try:
        return json.loads(llm_output)
    except Exception as e:
        st.error(f"Could not parse LLM output as JSON: {e}")
        return None

def add_comment(paragraph, comment_text):
    """
    Adds a simulated comment by appending an italicized bracketed note at the end of the paragraph.
    Real Word comments need deeper XML manipulation not covered here.
    """
    paragraph.add_run(f" [COMMENT: {comment_text}]").italic = True

# -----------------------------
# STREAMLIT UI
# -----------------------------
st.set_page_config(page_title="ADGM Corporate Agent", layout="centered")
st.title("ADGM Corporate Agent")
st.write("Upload your legal documents (.docx) for ADGM review and checklist validation.")

uploaded_files = st.file_uploader(
    "Select one or more .docx files",
    type="docx",
    accept_multiple_files=True
)

if uploaded_files:
    st.success(f"{len(uploaded_files)} document(s) uploaded.")
    
    detected_docs = []
    for uploaded_file in uploaded_files:
        try:
            doc = docx.Document(uploaded_file)
            doc_type = detect_document_type(doc)
            detected_docs.append(doc_type)
            st.write(f"📄 **{uploaded_file.name}** → Detected as: `{doc_type}`")

            search_query = f"ADGM rules for {doc_type}"
            references = retrieve_adgm_references(search_query, top_k=2)

            st.write("**Relevant ADGM References:**")
            for ref in references:
                st.markdown(f"- *{ref['source']}* — {ref['excerpt']}")

            doc_text = "\n".join([p.text for p in doc.paragraphs])
            llm_result = analyze_with_llm(doc_text, references)

            parsed_result = parse_llm_json(llm_result)
            if parsed_result:
                st.subheader("Identified Issues")
                st.json(parsed_result)

                # Prepare reviewed document copy for commenting
                reviewed_doc = docx.Document(uploaded_file)
                
                for issue in parsed_result.get("issues_found", []):
                    issue_section = issue.get("section", "")
                    suggestion = issue.get("suggestion", "")
                    # Simple paragraph matching on section text
                    for para in reviewed_doc.paragraphs:
                        if issue_section.lower() in para.text.lower():
                            add_comment(para, suggestion)
                            break

                # Save the reviewed docx to BytesIO
                reviewed_buffer = BytesIO()
                reviewed_doc.save(reviewed_buffer)
                reviewed_buffer.seek(0)

                # Save JSON report to BytesIO
                json_buffer = BytesIO()
                json_buffer.write(json.dumps(parsed_result, indent=2).encode())
                json_buffer.seek(0)

                # Provide download buttons
                st.download_button(
                    label="⬇ Download Reviewed DOCX",
                    data=reviewed_buffer,
                    file_name=f"reviewed_{uploaded_file.name}",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                st.download_button(
                    label="⬇ Download JSON Report",
                    data=json_buffer,
                    file_name=f"review_report_{uploaded_file.name}.json",
                    mime="application/json"
                )
            else:
                st.warning("No issues identified or unable to parse LLM output.")
                

        except Exception as e:
            st.error(f"Error reading {uploaded_file.name}: {e}")

    current_process = "Company Incorporation"
    required_docs = ADGM_CHECKLISTS[current_process]

    missing_docs = [req for req in required_docs if req not in detected_docs]
    st.subheader("Checklist Verification")
    st.write(f"Detected process: **{current_process}**")
    st.write(f"Required documents: {len(required_docs)}")
    st.write(f"Uploaded: {len(detected_docs)}")

    if missing_docs:
        st.warning(f"Missing required documents: {', '.join(missing_docs)}")
    else:
        st.success("All required documents appear to be present ✅")

else:
    st.info("Please upload documents to proceed.")
