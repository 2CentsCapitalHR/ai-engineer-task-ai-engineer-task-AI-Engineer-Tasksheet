"""Document parsing and type identification for ADGM documents."""

import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple
import logging
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from ..models import DocumentType, DocumentAnalysis
from ..config import ADGM_DOCUMENT_TYPES

logger = logging.getLogger(__name__)


class DocumentParser:
    """Parser for DOCX documents with ADGM-specific analysis."""
    
    def __init__(self):
        self.document_type_patterns = self._create_type_patterns()
    
    def parse_document(self, file_path: str) -> Dict:
        """Parse a DOCX document and extract relevant information."""
        try:
            doc = Document(file_path)
            
            # Extract basic information
            text_content = self._extract_text(doc)
            word_count = len(text_content.split())
            
            # Identify document type
            doc_type, confidence = self._identify_document_type(text_content)
            
            # Extract structured content
            structured_content = self._extract_structured_content(doc)
            
            # Extract metadata
            metadata = self._extract_metadata(doc)
            
            return {
                'filename': Path(file_path).name,
                'text_content': text_content,
                'word_count': word_count,
                'document_type': doc_type,
                'type_confidence': confidence,
                'structured_content': structured_content,
                'metadata': metadata,
                'docx_object': doc  # Keep for later modification
            }
            
        except Exception as e:
            logger.error(f"Failed to parse document {file_path}: {e}")
            raise
    
    def _extract_text(self, doc: Document) -> str:
        """Extract all text content from the document."""
        text_parts = []
        
        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text.strip())
        
        return '\n'.join(text_parts)
    
    def _identify_document_type(self, text_content: str) -> Tuple[DocumentType, float]:
        """Identify the type of document based on content analysis."""
        text_lower = text_content.lower()
        
        # Score each document type
        type_scores = {}
        
        for doc_type, patterns in self.document_type_patterns.items():
            score = 0
            total_patterns = len(patterns)
            
            for pattern in patterns:
                if re.search(pattern, text_lower):
                    score += 1
            
            # Calculate confidence as percentage of matched patterns
            confidence = (score / total_patterns) if total_patterns > 0 else 0
            type_scores[doc_type] = confidence
        
        # Find the best match
        if type_scores:
            best_type = max(type_scores, key=type_scores.get)
            best_confidence = type_scores[best_type]
            
            # Only return if confidence is above threshold
            if best_confidence >= 0.3:  # At least 30% of patterns matched
                return DocumentType(best_type), best_confidence
        
        return DocumentType.OTHER, 0.0
    
    def _create_type_patterns(self) -> Dict[str, List[str]]:
        """Create regex patterns for document type identification."""
        return {
            DocumentType.ARTICLES_OF_ASSOCIATION: [
                r'articles?\s+of\s+association',
                r'company\s+constitution',
                r'share\s+capital',
                r'directors?\s+powers',
                r'general\s+meetings?',
                r'dividend',
                r'winding\s+up'
            ],
            DocumentType.MEMORANDUM_OF_ASSOCIATION: [
                r'memorandum\s+of\s+association',
                r'objects?\s+of\s+the\s+company',
                r'liability\s+of\s+members',
                r'authorized\s+share\s+capital',
                r'company\s+name',
                r'registered\s+office'
            ],
            DocumentType.INCORPORATION_APPLICATION: [
                r'incorporation\s+application',
                r'application\s+for\s+registration',
                r'company\s+registration',
                r'proposed\s+company\s+name',
                r'nature\s+of\s+business',
                r'registered\s+address'
            ],
            DocumentType.UBO_DECLARATION: [
                r'ultimate\s+beneficial\s+owner',
                r'ubo\s+declaration',
                r'beneficial\s+ownership',
                r'controlling\s+interest',
                r'ownership\s+structure',
                r'25%\s+or\s+more'
            ],
            DocumentType.BOARD_RESOLUTION: [
                r'board\s+resolution',
                r'directors?\s+resolution',
                r'resolved\s+that',
                r'board\s+of\s+directors',
                r'meeting\s+of\s+directors',
                r'quorum\s+present'
            ],
            DocumentType.REGISTER_MEMBERS_DIRECTORS: [
                r'register\s+of\s+members',
                r'register\s+of\s+directors',
                r'shareholders?\s+register',
                r'directors?\s+register',
                r'member\s+details',
                r'director\s+details'
            ],
            DocumentType.SHAREHOLDER_RESOLUTION: [
                r'shareholders?\s+resolution',
                r'general\s+meeting',
                r'extraordinary\s+general\s+meeting',
                r'annual\s+general\s+meeting',
                r'special\s+resolution',
                r'ordinary\s+resolution'
            ],
            DocumentType.CHANGE_ADDRESS_NOTICE: [
                r'change\s+of\s+address',
                r'registered\s+office\s+address',
                r'new\s+address',
                r'address\s+change',
                r'relocation\s+notice',
                r'office\s+relocation'
            ],
            DocumentType.EMPLOYMENT_CONTRACT: [
                r'employment\s+contract',
                r'employment\s+agreement',
                r'terms\s+of\s+employment',
                r'job\s+description',
                r'salary',
                r'working\s+hours',
                r'notice\s+period'
            ],
            DocumentType.COMMERCIAL_AGREEMENT: [
                r'commercial\s+agreement',
                r'service\s+agreement',
                r'supply\s+agreement',
                r'partnership\s+agreement',
                r'terms\s+and\s+conditions',
                r'payment\s+terms'
            ],
            DocumentType.COMPLIANCE_POLICY: [
                r'compliance\s+policy',
                r'risk\s+management',
                r'data\s+protection',
                r'anti.money\s+laundering',
                r'know\s+your\s+customer',
                r'regulatory\s+compliance'
            ]
        }
    
    def _extract_structured_content(self, doc: Document) -> Dict:
        """Extract structured content like sections, clauses, etc."""
        structured = {
            'sections': [],
            'clauses': [],
            'tables': [],
            'signatures': []
        }
        
        # Extract sections (paragraphs with numbering or bold formatting)
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
            
            # Check for section headers (numbered or bold)
            if (re.match(r'^\d+\.', text) or 
                (paragraph.runs and any(run.bold for run in paragraph.runs))):
                structured['sections'].append({
                    'index': i,
                    'text': text,
                    'type': 'section_header'
                })
            
            # Check for clauses
            if re.match(r'^\d+\.\d+', text):
                structured['clauses'].append({
                    'index': i,
                    'text': text,
                    'clause_number': re.match(r'^(\d+\.\d+)', text).group(1)
                })
            
            # Check for signature lines
            if re.search(r'signature|signed|date.*signed', text.lower()):
                structured['signatures'].append({
                    'index': i,
                    'text': text
                })
        
        # Extract tables
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            
            structured['tables'].append({
                'index': i,
                'data': table_data,
                'rows': len(table.rows),
                'cols': len(table.rows[0].cells) if table.rows else 0
            })
        
        return structured
    
    def _extract_metadata(self, doc: Document) -> Dict:
        """Extract document metadata."""
        core_props = doc.core_properties
        
        return {
            'title': core_props.title or '',
            'author': core_props.author or '',
            'subject': core_props.subject or '',
            'created': core_props.created,
            'modified': core_props.modified,
            'last_modified_by': core_props.last_modified_by or '',
            'revision': core_props.revision,
            'category': core_props.category or ''
        }
    
    def add_comment_to_document(self, doc: Document, paragraph_index: int, 
                               comment_text: str, highlight_text: Optional[str] = None) -> Document:
        """Add a comment to a specific paragraph in the document."""
        try:
            if paragraph_index < len(doc.paragraphs):
                paragraph = doc.paragraphs[paragraph_index]
                
                # Add comment as a new paragraph after the target paragraph
                comment_paragraph = paragraph.insert_paragraph_after()
                comment_run = comment_paragraph.add_run(f"[ADGM COMMENT: {comment_text}]")
                
                # Style the comment
                comment_run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
                comment_run.italic = True
                
                # Highlight the original text if specified
                if highlight_text and highlight_text in paragraph.text:
                    for run in paragraph.runs:
                        if highlight_text in run.text:
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                
            return doc
            
        except Exception as e:
            logger.error(f"Failed to add comment: {e}")
            return doc
    
    def save_document_with_comments(self, doc: Document, output_path: str) -> str:
        """Save the document with added comments."""
        try:
            doc.save(output_path)
            logger.info(f"Document saved with comments: {output_path}")
            return output_path
        except Exception as e:
            logger.error(f"Failed to save document: {e}")
            raise
