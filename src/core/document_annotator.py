"""Document annotation system for inserting comments and highlights."""

import re
from typing import List, Dict, Optional, Tuple
import logging
from docx import Document
from docx.shared import RGBColor, Inches
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.shared import OxmlElement, qn
from ..models import DocumentIssue, CommentInsertion

logger = logging.getLogger(__name__)


class DocumentAnnotator:
    """Handles document annotation with comments and highlights."""
    
    def __init__(self):
        self.comment_styles = self._create_comment_styles()
    
    def annotate_document(self, doc: Document, issues: List[DocumentIssue], 
                         document_text: str) -> Document:
        """Annotate document with comments for identified issues."""
        
        try:
            # Create comment insertions from issues
            comment_insertions = self._create_comment_insertions(issues, document_text, doc)
            
            # Sort insertions by paragraph index (reverse order to avoid index shifting)
            comment_insertions.sort(key=lambda x: x.paragraph_index, reverse=True)
            
            # Insert comments
            for insertion in comment_insertions:
                self._insert_comment(doc, insertion)
            
            # Add summary at the beginning
            self._add_summary_section(doc, issues)
            
            logger.info(f"Added {len(comment_insertions)} comments to document")
            return doc
            
        except Exception as e:
            logger.error(f"Failed to annotate document: {e}")
            return doc
    
    def _create_comment_insertions(self, issues: List[DocumentIssue], 
                                  document_text: str, doc: Document) -> List[CommentInsertion]:
        """Create comment insertions from issues."""
        
        insertions = []
        paragraphs = [p.text for p in doc.paragraphs]
        
        for issue in issues:
            # Find the best paragraph to insert the comment
            paragraph_index = self._find_best_paragraph(issue, paragraphs)
            
            if paragraph_index is not None:
                comment_text = self._format_comment(issue)
                
                insertion = CommentInsertion(
                    paragraph_index=paragraph_index,
                    comment_text=comment_text,
                    highlight_text=self._extract_highlight_text(issue, paragraphs[paragraph_index]),
                    comment_type=self._get_comment_type(issue)
                )
                insertions.append(insertion)
        
        return insertions
    
    def _find_best_paragraph(self, issue: DocumentIssue, paragraphs: List[str]) -> Optional[int]:
        """Find the best paragraph to insert a comment for an issue."""
        
        # If issue has a specific section, try to find it
        if issue.section:
            section_keywords = issue.section.lower().split()
            
            for i, paragraph in enumerate(paragraphs):
                paragraph_lower = paragraph.lower()
                
                # Check if paragraph contains section keywords
                if any(keyword in paragraph_lower for keyword in section_keywords):
                    return i
        
        # Try to find paragraph based on issue content
        issue_keywords = self._extract_keywords_from_issue(issue)
        
        best_match_index = None
        best_match_score = 0
        
        for i, paragraph in enumerate(paragraphs):
            paragraph_lower = paragraph.lower()
            
            # Score paragraph based on keyword matches
            score = sum(1 for keyword in issue_keywords if keyword in paragraph_lower)
            
            if score > best_match_score:
                best_match_score = score
                best_match_index = i
        
        # Return best match if score is reasonable, otherwise return None
        return best_match_index if best_match_score > 0 else None
    
    def _extract_keywords_from_issue(self, issue: DocumentIssue) -> List[str]:
        """Extract keywords from issue description."""
        
        # Common legal terms and their variations
        keyword_map = {
            'jurisdiction': ['jurisdiction', 'court', 'adgm', 'dubai', 'abu dhabi'],
            'signature': ['signature', 'signed', 'signatory', 'execution'],
            'date': ['date', 'dated', 'day of'],
            'share capital': ['share', 'capital', 'shares', 'authorized'],
            'director': ['director', 'board', 'directors'],
            'member': ['member', 'shareholder', 'membership'],
            'company name': ['company', 'name', 'corporation'],
            'registered office': ['registered', 'office', 'address'],
            'objects': ['objects', 'purpose', 'business', 'activities']
        }
        
        keywords = []
        issue_text = issue.issue.lower()
        
        # Extract keywords based on issue content
        for category, terms in keyword_map.items():
            if any(term in issue_text for term in terms):
                keywords.extend(terms)
        
        # Add section-specific keywords
        if issue.section:
            keywords.extend(issue.section.lower().split())
        
        return list(set(keywords))  # Remove duplicates
    
    def _extract_highlight_text(self, issue: DocumentIssue, paragraph_text: str) -> Optional[str]:
        """Extract text to highlight based on the issue."""
        
        # Try to find specific text to highlight
        issue_keywords = self._extract_keywords_from_issue(issue)
        
        for keyword in issue_keywords:
            if keyword in paragraph_text.lower():
                # Find the actual text with proper case
                pattern = re.compile(re.escape(keyword), re.IGNORECASE)
                match = pattern.search(paragraph_text)
                if match:
                    return match.group()
        
        return None
    
    def _format_comment(self, issue: DocumentIssue) -> str:
        """Format a comment from an issue."""
        
        comment_parts = []
        
        # Add severity indicator
        severity_emoji = {
            'Low': '🟡',
            'Medium': '🟠', 
            'High': '🔴',
            'Critical': '🚨'
        }
        
        emoji = severity_emoji.get(issue.severity.value, '⚠️')
        comment_parts.append(f"{emoji} {issue.severity.value.upper()} ISSUE")
        
        # Add issue description
        comment_parts.append(f"Issue: {issue.issue}")
        
        # Add suggestion if available
        if issue.suggestion:
            comment_parts.append(f"Suggestion: {issue.suggestion}")
        
        # Add ADGM reference if available
        if issue.adgm_reference:
            comment_parts.append(f"Reference: {issue.adgm_reference}")
        
        return "\n".join(comment_parts)
    
    def _get_comment_type(self, issue: DocumentIssue) -> str:
        """Get comment type based on issue severity."""
        
        if issue.severity.value in ['High', 'Critical']:
            return 'critical'
        elif issue.severity.value == 'Medium':
            return 'warning'
        else:
            return 'info'
    
    def _insert_comment(self, doc: Document, insertion: CommentInsertion) -> None:
        """Insert a comment into the document."""
        
        try:
            if insertion.paragraph_index < len(doc.paragraphs):
                target_paragraph = doc.paragraphs[insertion.paragraph_index]
                
                # Highlight text if specified
                if insertion.highlight_text:
                    self._highlight_text_in_paragraph(target_paragraph, insertion.highlight_text)
                
                # Insert comment paragraph after the target paragraph
                comment_paragraph = self._insert_paragraph_after(target_paragraph)
                
                # Format comment based on type
                self._format_comment_paragraph(comment_paragraph, insertion)
                
        except Exception as e:
            logger.error(f"Failed to insert comment: {e}")
    
    def _highlight_text_in_paragraph(self, paragraph, highlight_text: str) -> None:
        """Highlight specific text in a paragraph."""
        
        try:
            paragraph_text = paragraph.text
            if highlight_text.lower() in paragraph_text.lower():
                
                # Find the text with proper case
                start_index = paragraph_text.lower().find(highlight_text.lower())
                if start_index != -1:
                    # Clear existing runs and recreate with highlighting
                    original_text = paragraph_text
                    paragraph.clear()
                    
                    # Add text before highlight
                    if start_index > 0:
                        paragraph.add_run(original_text[:start_index])
                    
                    # Add highlighted text
                    highlighted_run = paragraph.add_run(original_text[start_index:start_index + len(highlight_text)])
                    highlighted_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    
                    # Add text after highlight
                    if start_index + len(highlight_text) < len(original_text):
                        paragraph.add_run(original_text[start_index + len(highlight_text):])
                        
        except Exception as e:
            logger.error(f"Failed to highlight text: {e}")
    
    def _insert_paragraph_after(self, target_paragraph):
        """Insert a new paragraph after the target paragraph."""
        
        try:
            # Get the parent element
            parent = target_paragraph._element.getparent()
            
            # Create new paragraph element
            new_p = OxmlElement("w:p")
            
            # Insert after target paragraph
            parent.insert(parent.index(target_paragraph._element) + 1, new_p)
            
            # Create paragraph object
            from docx.text.paragraph import Paragraph
            new_paragraph = Paragraph(new_p, parent)
            
            return new_paragraph
            
        except Exception as e:
            logger.error(f"Failed to insert paragraph: {e}")
            # Fallback: add paragraph at the end
            return target_paragraph._parent.add_paragraph()
    
    def _format_comment_paragraph(self, paragraph, insertion: CommentInsertion) -> None:
        """Format a comment paragraph with appropriate styling."""
        
        try:
            # Add comment text
            run = paragraph.add_run(f"[ADGM REVIEW] {insertion.comment_text}")
            
            # Apply styling based on comment type
            if insertion.comment_type == 'critical':
                run.font.color.rgb = RGBColor(204, 0, 0)  # Dark red
                run.bold = True
            elif insertion.comment_type == 'warning':
                run.font.color.rgb = RGBColor(255, 102, 0)  # Orange
                run.bold = True
            else:
                run.font.color.rgb = RGBColor(0, 102, 204)  # Blue
            
            run.italic = True
            run.font.size = Inches(0.1)  # Slightly smaller font
            
            # Add border/background if possible
            paragraph.paragraph_format.left_indent = Inches(0.5)
            paragraph.paragraph_format.right_indent = Inches(0.5)
            
        except Exception as e:
            logger.error(f"Failed to format comment paragraph: {e}")
    
    def _add_summary_section(self, doc: Document, issues: List[DocumentIssue]) -> None:
        """Add a summary section at the beginning of the document."""
        
        try:
            # Insert summary at the beginning
            summary_paragraph = doc.paragraphs[0].insert_paragraph_before()
            
            # Add title
            title_run = summary_paragraph.add_run("ADGM COMPLIANCE REVIEW SUMMARY")
            title_run.bold = True
            title_run.font.size = Inches(0.15)
            title_run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
            
            # Add summary content
            summary_content = self._create_summary_content(issues)
            
            for line in summary_content:
                summary_para = summary_paragraph.insert_paragraph_after()
                summary_run = summary_para.add_run(line)
                summary_run.font.size = Inches(0.11)
            
            # Add separator
            separator_para = summary_paragraph.insert_paragraph_after()
            separator_run = separator_para.add_run("=" * 80)
            separator_run.font.color.rgb = RGBColor(128, 128, 128)
            
        except Exception as e:
            logger.error(f"Failed to add summary section: {e}")
    
    def _create_summary_content(self, issues: List[DocumentIssue]) -> List[str]:
        """Create summary content from issues."""
        
        if not issues:
            return ["✅ No compliance issues found. Document appears to be compliant with ADGM requirements."]
        
        # Count issues by severity
        severity_counts = {}
        for issue in issues:
            severity = issue.severity.value
            severity_counts[severity] = severity_counts.get(severity, 0) + 1
        
        summary_lines = [
            f"📊 Total Issues Found: {len(issues)}",
            ""
        ]
        
        # Add severity breakdown
        for severity in ['Critical', 'High', 'Medium', 'Low']:
            count = severity_counts.get(severity, 0)
            if count > 0:
                emoji = {'Critical': '🚨', 'High': '🔴', 'Medium': '🟠', 'Low': '🟡'}[severity]
                summary_lines.append(f"{emoji} {severity}: {count} issue(s)")
        
        summary_lines.extend([
            "",
            "📝 Key Issues:",
        ])
        
        # Add top issues
        for i, issue in enumerate(issues[:5]):  # Show top 5 issues
            summary_lines.append(f"{i+1}. {issue.issue}")
        
        if len(issues) > 5:
            summary_lines.append(f"... and {len(issues) - 5} more issue(s)")
        
        summary_lines.extend([
            "",
            "💡 Please review the detailed comments throughout the document for specific guidance.",
            ""
        ])
        
        return summary_lines
    
    def _create_comment_styles(self) -> Dict:
        """Create comment styling configurations."""
        
        return {
            'critical': {
                'color': RGBColor(204, 0, 0),  # Dark red
                'bold': True,
                'italic': True
            },
            'warning': {
                'color': RGBColor(255, 102, 0),  # Orange
                'bold': True,
                'italic': True
            },
            'info': {
                'color': RGBColor(0, 102, 204),  # Blue
                'bold': False,
                'italic': True
            }
        }
