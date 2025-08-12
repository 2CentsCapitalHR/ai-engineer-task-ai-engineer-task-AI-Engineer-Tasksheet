"""Create sample documents for demonstration purposes."""

from docx import Document
from pathlib import Path
import os


def create_sample_articles_of_association():
    """Create a sample Articles of Association document."""
    
    doc = Document()
    
    # Title
    title = doc.add_heading('ARTICLES OF ASSOCIATION', 0)
    title.alignment = 1  # Center alignment
    
    # Company details
    doc.add_heading('1. COMPANY NAME', level=1)
    doc.add_paragraph('The name of the company is "Sample Tech Solutions Limited".')
    
    doc.add_heading('2. REGISTERED OFFICE', level=1)
    doc.add_paragraph('The registered office of the company is situated in Abu Dhabi Global Market, United Arab Emirates.')
    
    doc.add_heading('3. OBJECTS OF THE COMPANY', level=1)
    doc.add_paragraph('The objects of the company are to carry on the business of technology consulting, software development, and related services.')
    
    doc.add_heading('4. SHARE CAPITAL', level=1)
    doc.add_paragraph('The authorized share capital of the company is AED 100,000 divided into 100,000 ordinary shares of AED 1 each.')
    
    doc.add_heading('5. LIABILITY OF MEMBERS', level=1)
    doc.add_paragraph('The liability of the members is limited by shares.')
    
    doc.add_heading('6. DIRECTORS POWERS', level=1)
    doc.add_paragraph('The directors may exercise all the powers of the company subject to the provisions of these Articles and applicable law.')
    
    doc.add_heading('7. GENERAL MEETINGS', level=1)
    doc.add_paragraph('An annual general meeting shall be held each year within 6 months of the financial year end.')
    
    doc.add_heading('8. GOVERNING LAW', level=1)
    # Intentional issue: Wrong jurisdiction
    doc.add_paragraph('These Articles shall be governed by the laws of Dubai Courts.')
    
    # Signature section - but incomplete
    doc.add_paragraph('\n\nSigned by:')
    doc.add_paragraph('Director: _________________')
    # Missing date field - another issue
    
    return doc


def create_sample_memorandum_of_association():
    """Create a sample Memorandum of Association document."""
    
    doc = Document()
    
    # Title
    title = doc.add_heading('MEMORANDUM OF ASSOCIATION', 0)
    title.alignment = 1
    
    doc.add_heading('1. COMPANY NAME', level=1)
    doc.add_paragraph('The name of the company is "Sample Tech Solutions Limited".')
    
    doc.add_heading('2. REGISTERED OFFICE', level=1)
    doc.add_paragraph('The registered office is situated in ADGM, UAE.')
    
    doc.add_heading('3. OBJECTS', level=1)
    doc.add_paragraph('The objects for which the company is established are:')
    doc.add_paragraph('a) To carry on business as technology consultants')
    doc.add_paragraph('b) To develop and market software solutions')
    doc.add_paragraph('c) To provide IT services and support')
    
    doc.add_heading('4. LIABILITY', level=1)
    doc.add_paragraph('The liability of members is limited by shares.')
    
    doc.add_heading('5. SHARE CAPITAL', level=1)
    doc.add_paragraph('The authorized share capital is AED 100,000.')
    
    # Missing some required sections - intentional for testing
    
    return doc


def create_sample_board_resolution():
    """Create a sample Board Resolution document."""
    
    doc = Document()
    
    # Title
    title = doc.add_heading('BOARD RESOLUTION', 0)
    title.alignment = 1
    
    doc.add_paragraph('Sample Tech Solutions Limited')
    doc.add_paragraph('ADGM Registration Number: ADGM-123456')
    doc.add_paragraph('')
    
    doc.add_paragraph('RESOLUTION OF THE BOARD OF DIRECTORS')
    doc.add_paragraph('Passed on [DATE TO BE FILLED]')
    doc.add_paragraph('')
    
    doc.add_paragraph('At a meeting of the Board of Directors of the Company held on [DATE], the following resolutions were passed:')
    doc.add_paragraph('')
    
    doc.add_paragraph('RESOLVED THAT:')
    doc.add_paragraph('')
    
    doc.add_paragraph('1. The company shall open a corporate bank account with a licensed bank in ADGM.')
    doc.add_paragraph('')
    
    doc.add_paragraph('2. The directors are hereby authorized to execute all necessary documents for the bank account opening.')
    doc.add_paragraph('')
    
    doc.add_paragraph('3. The company secretary is authorized to file this resolution with the ADGM Registration Authority.')
    doc.add_paragraph('')
    
    # Signature section
    doc.add_paragraph('Quorum present: 2 directors')
    doc.add_paragraph('')
    doc.add_paragraph('Chairman: _________________')
    doc.add_paragraph('Date: _________________')
    doc.add_paragraph('')
    doc.add_paragraph('Director: _________________')
    doc.add_paragraph('Date: _________________')
    
    return doc


def create_sample_employment_contract():
    """Create a sample Employment Contract document."""
    
    doc = Document()
    
    # Title
    title = doc.add_heading('EMPLOYMENT CONTRACT', 0)
    title.alignment = 1
    
    doc.add_paragraph('This Employment Agreement is entered into between:')
    doc.add_paragraph('')
    
    doc.add_paragraph('EMPLOYER: Sample Tech Solutions Limited')
    doc.add_paragraph('ADGM Registration: ADGM-123456')
    doc.add_paragraph('Address: ADGM, Abu Dhabi, UAE')
    doc.add_paragraph('')
    
    doc.add_paragraph('EMPLOYEE: [Employee Name]')
    doc.add_paragraph('Passport Number: [Number]')
    doc.add_paragraph('')
    
    doc.add_heading('1. POSITION AND DUTIES', level=1)
    doc.add_paragraph('The Employee shall serve as Software Developer and perform duties as assigned.')
    
    doc.add_heading('2. SALARY', level=1)
    doc.add_paragraph('The Employee shall receive a monthly salary of AED 15,000.')
    
    doc.add_heading('3. WORKING HOURS', level=1)
    doc.add_paragraph('Normal working hours are 40 hours per week, Sunday to Thursday.')
    
    doc.add_heading('4. NOTICE PERIOD', level=1)
    doc.add_paragraph('Either party may terminate this agreement with 30 days written notice.')
    
    doc.add_heading('5. GOVERNING LAW', level=1)
    doc.add_paragraph('This contract shall be governed by ADGM Employment Regulations.')
    
    # Signature section
    doc.add_paragraph('')
    doc.add_paragraph('EMPLOYER:')
    doc.add_paragraph('Signature: _________________')
    doc.add_paragraph('Name: _________________')
    doc.add_paragraph('Date: _________________')
    doc.add_paragraph('')
    doc.add_paragraph('EMPLOYEE:')
    doc.add_paragraph('Signature: _________________')
    doc.add_paragraph('Name: _________________')
    doc.add_paragraph('Date: _________________')
    
    return doc


def create_sample_documents():
    """Create all sample documents."""
    
    # Create demo directory
    demo_dir = Path("demo/sample_documents")
    demo_dir.mkdir(parents=True, exist_ok=True)
    
    # Create documents
    documents = {
        "Sample_Articles_of_Association.docx": create_sample_articles_of_association(),
        "Sample_Memorandum_of_Association.docx": create_sample_memorandum_of_association(),
        "Sample_Board_Resolution.docx": create_sample_board_resolution(),
        "Sample_Employment_Contract.docx": create_sample_employment_contract()
    }
    
    # Save documents
    for filename, doc in documents.items():
        file_path = demo_dir / filename
        doc.save(str(file_path))
        print(f"Created: {file_path}")
    
    print(f"\nSample documents created in: {demo_dir}")
    print("\nThese documents contain intentional issues for demonstration:")
    print("- Articles of Association: Wrong jurisdiction reference")
    print("- Memorandum of Association: Missing required sections")
    print("- Board Resolution: Missing date fields")
    print("- Employment Contract: Generally compliant")


if __name__ == "__main__":
    create_sample_documents()
